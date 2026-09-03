# -*- coding: utf-8 -*-
"""Los comentarios del cierre: la columna, la nota del desplegable y el Word.

Owner, 2026-09-03: *«hay una celda al final del P&L que dice Commentary pero no
tiene forma para que sea editable»* y, enseguida: *«una vez que abrimos el popup,
abajo de los números abrir un box para editar comentarios… que se guarde con la
versión del mes, con el mes. Y si cambio de mes, que se me abra otra opción para
agregar notas. Y una vez que se impriman en Word, estas notas aparezcan en el box
editable»*.
"""
import inspect
from pathlib import Path

from docx import Document

from app.api import comentario_pl_api as api
from app.export.cierre_word import build_cierre_docx

FRONT = Path(__file__).resolve().parents[2] / "frontend"
CIERRE = FRONT / "app/month-end/pl"


def test_el_comentario_se_GUARDA():
    """⚠️ Uno que se pierde al recargar es peor que ninguno: el que lo escribió
    cree que quedó."""
    assert "annotations" in inspect.getsource(api).lower()


def test_la_llave_lleva_el_MES():
    """Owner: «que se guarde con el mes; si cambio de mes, que se me abra otra
    opción».

    ⚠️ Y además es lo correcto: la explicación de julio no explica agosto, y una
    nota sin mes se arrastraría a todos los cierres siguientes diciendo algo que
    ya no es cierto.
    """
    fuente = inspect.getsource(api.guardar)
    assert "Annotation.month == cuerpo.mes" in fuente
    pantalla = (CIERRE / "DetalleCelda.tsx").read_text(encoding="utf-8")
    assert "[escNota, mes, ref]" in pantalla, (
        "la nota del desplegable dejó de recargarse al cambiar de mes")


def test_guardar_es_UPSERT_y_no_deja_duplicados():
    """⚠️ `POST /annotations/` siempre CREA. Para una celda que se edita en el
    lugar, eso deja una fila por cada corrección de una coma — y la que gane es
    la que la consulta devuelva primero, o sea al azar."""
    fuente = inspect.getsource(api.guardar)
    assert "principal.body = texto" in fuente
    assert "for a in sobrantes:" in fuente


def test_vaciar_la_celda_BORRA():
    """Guardar una cadena vacía dejaría una fila fantasma que mañana alguien
    cuenta como «hay comentario»."""
    fuente = inspect.getsource(api.guardar)
    assert "if not texto:" in fuente and "await db.delete(a)" in fuente


def test_no_se_agrego_pl_a_las_SECCIONES_de_asignaciones():
    """⚠️ `SECTIONS` es el vocabulario de quién es responsable de qué. Sumarle
    «pl» le inventaría a la pantalla de colaboración una sección más, con su
    responsable y su estado, que nadie pidió."""
    from app.models.section_assignment import SECTIONS
    assert "pl" not in SECTIONS


def test_la_nota_de_una_CELDA_no_pisa_la_de_al_lado():
    """`payroll` de Rooms y `opex` de Rooms son dos celdas distintas."""
    pantalla = (CIERRE / "DetalleCelda.tsx").read_text(encoding="utf-8")
    assert "`celda:${c.clase}:${c.clave" in pantalla


def test_se_guarda_al_SALIR_del_campo_y_no_por_tecla():
    """Por tecla serían treinta llamadas por nota, y una carrera donde gana la
    que conteste última — que no es la última que se escribió."""
    pantalla = (CIERRE / "DetalleCelda.tsx").read_text(encoding="utf-8")
    assert "onBlur={guardarNota}" in pantalla
    pagina = (CIERRE / "page.tsx").read_text(encoding="utf-8")
    assert "onBlur={e => guardarComentario(" in pagina


def test_lo_escrito_NO_se_pierde_si_falla_la_red():
    """Perderlo sería lo peor que puede hacer una nota: queda en el campo y se
    reintenta al volver a salir de él."""
    pantalla = (CIERRE / "DetalleCelda.tsx").read_text(encoding="utf-8")
    cuerpo = pantalla[pantalla.index("const guardarNota"):]
    assert "} catch {" in cuerpo[:900]


def test_las_notas_se_IMPRIMEN_dentro_del_recuadro(tmp_path):
    """Owner: «que estas notas aparezcan en el box editable».

    ⚠️ Dentro y no en un párrafo aparte: el recuadro es el lugar donde se
    comenta, y separar lo escrito de donde se escribe haría que en la reunión se
    comente dos veces lo mismo.
    """
    cuadro = {
        "titulo": "X", "subtitulo": "y",
        "columnas": [{"label": "A", "ancho": 20, "formato": "texto"},
                     {"label": "B", "ancho": 10, "formato": "usd2"}],
        "filas": [{"label": "uno", "es_total": False, "valores": [1.5]}],
        "comentarios": ["0110 · Rooms — subió por dos plazas nuevas"],
    }
    ruta = tmp_path / "x.docx"
    ruta.write_bytes(build_cierre_docx([cuadro], "H", "C", "Julio 2026", "v"))
    caja = Document(str(ruta)).tables[-1]
    textos = [p.text.strip() for p in caja.rows[0].cells[0].paragraphs]
    assert any("subió por dos plazas nuevas" in t for t in textos)
    # Y queda espacio para escribir a mano ENCIMA de lo que ya está.
    assert textos.count("") >= 3


def test_un_cuadro_SIN_notas_sigue_trayendo_su_recuadro_vacio(tmp_path):
    """El recuadro es para comentar en la reunión: sacarlo cuando no hay nada
    escrito dejaría sin dónde escribir justo al que no escribió antes."""
    cuadro = {"titulo": "X", "subtitulo": "y",
              "columnas": [{"label": "A", "ancho": 20, "formato": "texto"}],
              "filas": [{"label": "uno", "es_total": False, "valores": []}]}
    ruta = tmp_path / "y.docx"
    ruta.write_bytes(build_cierre_docx([cuadro], "H", "C", "Julio 2026", "v"))
    caja = Document(str(ruta)).tables[-1]
    assert len(caja.rows[0].cells[0].paragraphs) >= 4


def test_cada_nota_sale_con_el_RENGLON_al_que_pertenece():
    """Una lista de frases sueltas dentro de un recuadro no dice de qué habla
    cada una."""
    pagina = (CIERRE / "page.tsx").read_text(encoding="utf-8")
    assert "function notasDe(" in pagina
    assert "— ${texto.trim()}" in pagina


def test_el_comentario_BAJA_con_el_Excel():
    """Escribir la explicación y que el Excel salga con la columna vacía es el
    defecto que el owner ya reportó una vez: «el excel no baja lo que está
    viendo» (2026-08-27)."""
    pagina = (CIERRE / "page.tsx").read_text(encoding="utf-8")
    assert "comentarios[f.code] ?? null," in pagina


def test_NINGUN_punto_puede_olvidarse_el_mes():
    """Owner, 2026-09-03: «que las notas queden ligadas al mes actual, para que
    viajen con el mes actual».

    ⚠️ La llave es (escenario, "pl", renglón, mes) y el mes NO tiene valor por
    defecto: una lectura sin mes traería las notas de todos los cierres
    mezcladas, y una escritura sin mes las dejaría pegadas al mes 0 —donde
    ninguna pantalla las volvería a encontrar—.

    Se cuentan los puntos de llamada: son cinco —la columna del P&L, la nota
    del desplegable, sus dos guardados y el Word— y los cinco tienen que pasar
    el mes. Uno solo que se lo olvide devuelve el defecto en una pantalla, que
    es como no se nota.
    """
    puntos = 0
    for archivo in ("page.tsx", "DetalleCelda.tsx"):
        src = (CIERRE / archivo).read_text(encoding="utf-8")
        for linea in src.splitlines():
            if "getComentariosPL(" in linea or "guardarComentarioPL(" in linea:
                puntos += 1
                assert "mes" in linea, f"sin el mes: {archivo} → {linea.strip()}"
    assert puntos >= 5, f"quedaron {puntos} puntos de llamada, se esperaban 5"


def test_el_backend_EXIGE_el_mes_al_leer():
    """Sin valor por defecto: una lectura sin mes es un error, no «todos»."""
    fuente = inspect.getsource(api.listar)
    assert "mes: int = Query(..., ge=1, le=12)" in fuente, (
        "el mes dejó de ser obligatorio al leer: las notas de todos los "
        "cierres saldrían mezcladas")


def test_las_dos_consultas_del_backend_filtran_por_MES():
    """Leer y guardar. Que una sola se lo saltee alcanza para que una nota de
    julio aparezca en agosto o para que se pisen entre meses."""
    assert "Annotation.month == mes" in inspect.getsource(api.listar)
    assert "Annotation.month == cuerpo.mes" in inspect.getsource(api.guardar)


def test_hay_boton_de_guardar_Y_guardado_automatico():
    """Owner, 2026-09-03: «y que guarda… no sé si se necesita un botón de
    guardar ahí mismo».

    Las dos cosas, y no una:

    * el guardado al salir del campo es lo que impide PERDER lo escrito —
      guardar sólo con el botón haría que salirse sin tocarlo pierda la nota;
    * el botón es lo que quita la DUDA, que era el problema real. En una
      reunión, escribir la explicación y quedarse sin saber si quedó es tan
      malo como que no quede.
    """
    pantalla = (CIERRE / "DetalleCelda.tsx").read_text(encoding="utf-8")
    assert "onBlur={guardarNota}" in pantalla
    assert "onClick={guardarNota}" in pantalla


def test_el_estado_de_la_nota_se_DICE_siempre():
    """No sólo mientras guarda: «sin guardar», «guardando…» o «guardada»."""
    pantalla = (CIERRE / "DetalleCelda.tsx").read_text(encoding="utf-8")
    for estado in ('"sin guardar"', '"guardando…"', '"guardada"'):
        assert estado in pantalla, f"falta el estado {estado}"


def test_pendiente_se_deduce_comparando_con_lo_GUARDADO():
    """⚠️ Y no con una bandera aparte. El botón le roba el foco al campo y
    dispara el `onBlur`; con una bandera, esa carrera guarda dos veces o deja
    el botón diciendo «Guardar» sobre algo ya guardado."""
    pantalla = (CIERRE / "DetalleCelda.tsx").read_text(encoding="utf-8")
    assert "nota.trim() !== guardado.trim()" in pantalla
    assert "if (texto === guardado.trim()) return;" in pantalla


def test_si_falla_la_red_el_boton_SIGUE_ofreciendo_guardar():
    """Marcarla como guardada cuando el guardado falló es la peor mentira que
    puede decir esta pantalla."""
    pantalla = (CIERRE / "DetalleCelda.tsx").read_text(encoding="utf-8")
    cuerpo = pantalla[pantalla.index("const guardarNota"):]
    cuerpo = cuerpo[:cuerpo.index("}, [escNota")]
    # `setGuardado` sólo dentro del `try`, después de que el servidor contestó.
    assert cuerpo.index("setGuardado(texto)") < cuerpo.index("} catch {")
