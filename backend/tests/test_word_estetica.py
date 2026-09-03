# -*- coding: utf-8 -*-
"""Cómo se ve el Word del cierre.

Owner, 2026-09-03, con el documento de julio delante: *«hay cuadros que no
tienen datos, y los cuadros deben ser más pequeños, y que se vean más estéticos
y profesional»*.
"""
from docx import Document
from docx.oxml.ns import qn

from app.export.cierre_word import build_cierre_docx

CUADRO = {
    "titulo": "Prueba",
    "subtitulo": "sub",
    "columnas": [
        {"label": "Line Item", "ancho": 26, "formato": "texto"},
        {"label": "ACTUAL", "ancho": 17, "formato": "usd2"},
        {"label": "BUDGET", "ancho": 17, "formato": "usd2"},
    ],
    "filas": [
        {"label": "Rooms", "es_total": False, "valores": [36218.36, 46500.0]},
        {"label": "F&B", "es_total": False, "valores": [2433.47, 0.0]},
        {"label": "Total Revenue", "es_total": True, "valores": [56001.80, 66922.0]},
    ],
}


def _doc(tmp_path, cuadros=None):
    b = build_cierre_docx(cuadros or [CUADRO], "Hotel", "Cierre", "Julio 2026", "v")
    ruta = tmp_path / "x.docx"
    ruta.write_bytes(b)
    return Document(str(ruta))


def test_la_tabla_NO_lleva_rejilla(tmp_path):
    """⚠️ Antes usaba el estilo `Table Grid`, que dibuja una caja alrededor de
    cada celda.

    Un estado financiero impreso no lleva rejilla: lleva reglas horizontales.
    La rejilla completa hace que cada celda pese lo mismo, que es lo contrario
    de lo que un estado necesita — ahí el ojo tiene que caer en los totales.
    """
    tb = _doc(tmp_path).tables[0]
    bordes = tb._tbl.tblPr.find(qn("w:tblBorders"))
    assert bordes is not None, "volvió el estilo con rejilla"
    for el in bordes:
        assert el.get(qn("w:val")) == "none", (
            f"la tabla volvió a dibujar el borde «{el.tag.split('}')[1]}»")


def test_los_TOTALES_llevan_regla_arriba(tmp_path):
    """La convención del estado impreso: la línea dice «acá se cierra algo»."""
    tb = _doc(tmp_path).tables[0]
    total = tb.rows[3]                      # 0 = encabezado
    assert "Total Revenue" in total.cells[0].text
    tcb = total.cells[0]._tc.tcPr.find(qn("w:tcBorders"))
    assert tcb is not None and tcb.find(qn("w:top")) is not None


def test_una_fila_normal_NO_lleva_regla(tmp_path):
    """Una regla bajo cada fila devuelve la rejilla que se acaba de sacar."""
    tb = _doc(tmp_path).tables[0]
    c = tb.rows[1].cells[0]
    tcp = c._tc.tcPr
    tcb = tcp.find(qn("w:tcBorders")) if tcp is not None else None
    assert tcb is None or len(tcb) == 0


def test_las_celdas_van_APRETADAS(tmp_path):
    """⚠️ Word deja 108 twips a cada lado por defecto: con ocho columnas son
    casi cinco centímetros de aire. Apretarlos es lo que de verdad achica el
    cuadro; bajar la letra sin esto sólo lo hace ilegible."""
    tb = _doc(tmp_path).tables[0]
    mar = tb._tbl.tblPr.find(qn("w:tblCellMar"))
    assert mar is not None, "los márgenes de celda volvieron al default"
    for el in mar:
        assert int(el.get(qn("w:w"))) <= 60


def test_no_queda_ningun_run_VACIO_de_10_puntos(tmp_path):
    """⚠️ `celda.text = ""` no vacía: deja un run vacío que hereda el estilo
    Normal —10 pt— y ese run invisible fija la altura de la fila. Con letra de
    7,5 pt son casi tres milímetros por fila que nadie pidió."""
    tb = _doc(tmp_path).tables[0]
    for fila in tb.rows:
        for celda in fila.cells:
            for r in celda.paragraphs[0].runs:
                assert r.text != "", "quedó un run vacío inflando la fila"
                assert r.font.size is not None, "un run sin tamaño hereda 10 pt"


def test_el_ancho_lo_deciden_las_COLUMNAS_y_no_el_texto(tmp_path):
    """⚠️ Antes iba con `autofit` y Word repartía por el CONTENIDO: una columna
    con «TOTAL RENT AND MANAGEMENT FEES» se comía el espacio de los números.
    Los `ancho` ya viajaban en el payload y no los miraba nadie."""
    d = _doc(tmp_path)
    tb = d.tables[0]
    anchos = [c.width for c in tb.rows[0].cells]
    assert all(a is not None for a in anchos), "las columnas no tienen ancho"
    # La primera pesa 26 contra 17: tiene que ser más ancha.
    assert anchos[0] > anchos[1]
    # Y todo junto entra en la página.
    sec = d.sections[-1]
    util = sec.page_width - sec.left_margin - sec.right_margin
    assert sum(a for a in anchos) <= util + 5000


def test_la_letra_es_de_SIETE_Y_MEDIO(tmp_path):
    """Más chica que los 8,5 anteriores. El pedido fue «más pequeños»."""
    tb = _doc(tmp_path).tables[0]
    for fila in tb.rows:
        for celda in fila.cells:
            for r in celda.paragraphs[0].runs:
                assert r.font.size.pt <= 7.5
