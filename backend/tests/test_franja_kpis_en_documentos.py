# -*- coding: utf-8 -*-
"""Las estadísticas viajan con cada cuadro, en Word y en Excel.

Owner, 2026-09-03: *«no están saliendo las estadísticas en cada tab»*.

⚠️ En la PANTALLA la franja se dibuja una sola vez, arriba de los sub-tabs, así
que se ve en todos. En un documento **cada hoja se lee sola** —se imprime, se
manda suelta, se pega en un correo— y sin las estadísticas al lado los montos no
tienen contra qué leerse: 56.001 de ingreso con 132 noches vendidas dice algo muy
distinto que con 400.
"""
import io as _io
from pathlib import Path

import openpyxl
from docx import Document

from app.export.cierre_word import build_cierre_docx
from app.export.cuadro_excel import build_cuadros_workbook

FRONT = Path(__file__).resolve().parents[2] / "frontend"
CIERRE = FRONT / "app/month-end/pl"

KPIS = {
    "kpis_columnas": ["ACTUAL Final 2026", "BUDGET Final 2026"],
    "kpis": [
        {"label": "Total Rooms Occupied", "valores": [132, 124]},
        {"label": "% Occupancy", "valores": [0.2661, 0.25]},
        {"label": "Average Daily Room Only", "valores": [255.44, 375.0]},
    ],
}
CUADRO = {
    "titulo": "Payroll x Depto", "subtitulo": "Julio 2026 · USD", **KPIS,
    "columnas": [{"label": "Depto", "ancho": 26, "formato": "texto"},
                 {"label": "ACTUAL", "ancho": 16, "formato": "usd2"}],
    "filas": [{"label": "0110 · Rooms", "es_total": False, "valores": [9518.26]}],
}


def test_el_EXCEL_trae_la_franja_arriba_del_cuadro():
    wb = openpyxl.load_workbook(_io.BytesIO(build_cuadros_workbook([CUADRO])))
    ws = wb.worksheets[0]
    textos = [ws.cell(r, 1).value for r in range(1, 12)]
    assert "ESTADÍSTICAS" in textos
    assert "Total Rooms Occupied" in textos


def test_la_cabecera_del_cuadro_NO_queda_pisada_por_la_franja():
    """⚠️ Las filas del cuadro eran constantes fijas. Con la franja delante,
    escribir la tabla en la fila 4 la escribiría ENCIMA de las estadísticas —y
    el cuadro seguiría viéndose bien, sólo que sin ellas."""
    wb = openpyxl.load_workbook(_io.BytesIO(build_cuadros_workbook([CUADRO])))
    ws = wb.worksheets[0]
    col1 = [ws.cell(r, 1).value for r in range(1, 15)]
    assert col1.index("Depto") > col1.index("Average Daily Room Only")


def test_sin_franja_el_cuadro_arranca_donde_siempre():
    """Un cuadro sin estadísticas no puede quedar con cuatro filas en blanco
    arriba."""
    sin = {k: v for k, v in CUADRO.items() if k not in ("kpis", "kpis_columnas")}
    wb = openpyxl.load_workbook(_io.BytesIO(build_cuadros_workbook([sin])))
    ws = wb.worksheets[0]
    assert ws.cell(4, 1).value == "Depto"


def test_el_WORD_trae_la_franja_antes_de_la_tabla(tmp_path):
    ruta = tmp_path / "x.docx"
    ruta.write_bytes(build_cierre_docx([CUADRO], "H", "C", "Julio 2026", "v"))
    d = Document(str(ruta))
    primera = d.tables[0]
    assert "ESTADÍSTICAS" in primera.rows[0].cells[0].text
    # Y la tabla del cuadro viene DESPUÉS.
    assert any("0110 · Rooms" in c.text
               for t in d.tables[1:] for r in t.rows for c in r.cells)


def test_la_ocupacion_sale_como_PORCENTAJE_y_la_tarifa_como_dolares(tmp_path):
    """El formato lo decide el rótulo. Sin eso, 0,2661 de ocupación saldría
    como «0» y la tarifa como «255»."""
    ruta = tmp_path / "y.docx"
    ruta.write_bytes(build_cierre_docx([CUADRO], "H", "C", "Julio 2026", "v"))
    texto = "\n".join(c.text for r in Document(str(ruta)).tables[0].rows
                      for c in r.cells)
    assert "26.6%" in texto
    assert "255.44" in texto


def test_el_MODELO_del_endpoint_conserva_la_franja():
    """⚠️ Igual que con las notas: si el modelo no la declara, Pydantic la
    descarta en silencio y el documento sale sin ella sin que nada falle."""
    from app.api.export_api import Cuadro
    c = Cuadro(**CUADRO)
    assert len(c.kpis) == 3
    assert c.kpis_columnas == ["ACTUAL Final 2026", "BUDGET Final 2026"]
    assert "kpis" in c.model_dump()


def test_los_ROTULOS_son_los_mismos_que_en_pantalla():
    """Que el documento llame «Average Daily Room Only» a lo que la pantalla
    llama otra cosa obliga a comprobar que son el mismo número."""
    pagina = (CIERRE / "page.tsx").read_text(encoding="utf-8")
    pantalla = (CIERRE / "Estadisticas.tsx").read_text(encoding="utf-8")
    for rotulo in ("Total available Rooms", "Total Rooms Occupied",
                   "% Occupancy", "Average Daily Room Only", "Total RevPAR"):
        assert rotulo in pagina, f"el documento no trae «{rotulo}»"
        assert rotulo in pantalla, f"la pantalla ya no dice «{rotulo}»"


def test_si_las_estadisticas_fallan_el_documento_SALE_igual():
    """Sin franja se pierde contexto; sin documento se pierde la reunión."""
    pagina = (CIERRE / "page.tsx").read_text(encoding="utf-8")
    cuerpo = pagina[pagina.index("async function franjaKpis()"):]
    cuerpo = cuerpo[:cuerpo.index("async function bajarExcel")]
    assert ".catch(() => null)" in cuerpo
    assert "return null;" in cuerpo
