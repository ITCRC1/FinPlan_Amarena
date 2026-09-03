# -*- coding: utf-8 -*-
"""El reporte de cierre en Word, con espacio para comentar cada cuadro.

Owner, 2026-09-02: *«necesito hacer un documento Word con todos los tabs
activos; la idea es hacer un reporte con comentarios, así que dejá espacio entre
los tabs para poder comentar. Los reportes y cortes deben ser muy profesionales,
debe verse súper bien. Y siempre deben salir los tabs que estén activos en la
vista»*.

## La pantalla manda lo que está viendo

Mismo contrato que el Excel (`app/export/cuadro_excel.py`): llegan `cuadros` ya
armados —columnas, filas, valores— y acá sólo se les da forma. **No se
recalcula nada.** Buscar los datos desde el servidor obligaría a reimplementar
cada pantalla y arriesgaría que el documento diga algo distinto de lo que el
owner tenía delante cuando decidió exportarlo.

De ahí sale gratis lo de *«siempre deben salir los tabs que estén activos»*: la
pantalla manda los que están activos, porque son los únicos que dibuja.

## Vertical, y las páginas anchas en horizontal

Owner, 2026-09-02, eligiendo entre tres formatos. El documento se lee como
reporte —vertical— pero un cuadro de catorce columnas ahí no entra: el `P&L
Statement` lleva mes y YTD con seis columnas cada uno. Los anchos abren su
propia sección apaisada y vuelven a vertical después.

⚠️ **El umbral se mide en columnas, no en caracteres.** Una tabla de Word
reparte el ancho disponible entre sus columnas, así que lo que decide si se lee
o no es cuántas hay, no cuán largo es el texto — que Word parte en varias líneas
sin romper nada.

## El espacio para comentar

Cada cuadro termina en un bloque **Comentarios** con renglones vacíos. Es una
tabla de una celda y no párrafos sueltos a propósito: el marco se ve al
imprimir, sobrevive a que alguien escriba de más —crece— y no se confunde con el
cuadro de arriba.
"""
from __future__ import annotations

import io
from datetime import date

from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Emu, Pt, RGBColor

DOCX = ("application/vnd.openxmlformats-officedocument"
        ".wordprocessingml.document")

#: A partir de acá el cuadro va apaisado. Con siete columnas en vertical cada
#: una queda en ~2,3 cm, que todavía admite un número de nueve dígitos; con
#: ocho ya se parte.
COLUMNAS_APAISADO = 8

#: Cuántos renglones deja el bloque de comentarios. Cuatro llenan un tercio de
#: página sin empujar el cuadro siguiente a una hoja de más.
RENGLONES_COMENTARIO = 4

GRIS = RGBColor(0x60, 0x66, 0x6E)
NEGRO = RGBColor(0x1A, 0x1D, 0x21)
ROJO = RGBColor(0xB3, 0x26, 0x1E)
MARCA = RGBColor(0x1B, 0x3A, 0x5C)


# ─── Formato de números ───────────────────────────────────────────────────────
def _texto(valor, formato: str) -> str:
    """El valor tal como se ve en pantalla.

    ⚠️ Acá **sí** se formatea a texto, al revés que en el Excel. Un `.xlsx` con
    «$1,234.00» como cadena no se puede sumar ni graficar, y por eso allá viajan
    números; un Word no se recalcula nunca — lo que se quiere es que el número
    se lea igual que en la pantalla de la que salió.
    """
    if valor is None or valor == "":
        return ""
    if isinstance(valor, str):
        return valor
    if formato == "pct":
        return f"({abs(valor) * 100:,.1f}%)" if valor < 0 else f"{valor * 100:,.1f}%"
    if formato in ("num", "num1"):
        dec = 1 if formato == "num1" else 0
        return f"({abs(valor):,.{dec}f})" if valor < 0 else f"{valor:,.{dec}f}"
    dec = 2 if formato == "usd2" else 0
    return f"({abs(valor):,.{dec}f})" if valor < 0 else f"{valor:,.{dec}f}"


def _negativo(valor) -> bool:
    return isinstance(valor, (int, float)) and valor < 0


# ─── Piezas de Word ───────────────────────────────────────────────────────────
def _sombra(celda, hex_color: str) -> None:
    """Fondo de celda. `python-docx` no lo expone; se arma el XML a mano."""
    tc = celda._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hex_color)
    tc.append(shd)



def _vaciar(celda):
    """Deja la celda con UN párrafo y sin runs, y lo devuelve.

    ⚠️ `celda.text = ""` no vacía: deja un run vacío que hereda el estilo
    Normal —10 pt—, y ese run invisible fija la altura de la fila. Con letra de
    7,5 pt eso son casi tres milímetros de aire por fila que nadie pidió, y en
    un cuadro de sesenta filas se nota. Owner, 2026-09-03: «los cuadros deben
    ser más pequeños».
    """
    p = celda.paragraphs[0]
    for r in list(p.runs):
        r._r.getparent().remove(r._r)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    return p


def _sin_rejilla(tabla) -> None:
    """Le quita TODOS los bordes a la tabla.

    ⚠️ La versión anterior usaba el estilo `Table Grid`, que dibuja una caja
    negra alrededor de cada celda. Owner, 2026-09-03: *«los cuadros deben ser
    más pequeños, y que se vean más estéticos y profesional»*.

    Un estado financiero impreso NO lleva rejilla: lleva **reglas
    horizontales** —una bajo el encabezado, una sobre cada total— y ninguna
    vertical. La rejilla completa hace que cada celda pese lo mismo, que es lo
    contrario de lo que un estado necesita: ahí el ojo tiene que caer en los
    totales.
    """
    tbl_pr = tabla._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for lado in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{lado}")
        el.set(qn("w:val"), "none")
        el.set(qn("w:sz"), "0")
        borders.append(el)
    tbl_pr.append(borders)


def _regla(celda, lado: str, grosor: int, color: str) -> None:
    """Una regla horizontal en una celda. `grosor` va en octavos de punto."""
    tc_pr = celda._tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    el = OxmlElement(f"w:{lado}")
    el.set(qn("w:val"), "single")
    el.set(qn("w:sz"), str(grosor))
    el.set(qn("w:color"), color)
    borders.append(el)


def _apretar(tabla, arriba: int = 14, abajo: int = 14,
             izq: int = 60, der: int = 60) -> None:
    """Márgenes internos de celda, en vigésimos de punto (twips).

    Word deja 108 twips a cada lado por defecto —más de medio centímetro por
    columna—, y con ocho columnas eso solo son casi cinco centímetros de aire.
    Apretarlos es lo que de verdad achica el cuadro; bajar la letra sin esto
    solo lo hace ilegible.
    """
    tbl_pr = tabla._tbl.tblPr
    mar = OxmlElement("w:tblCellMar")
    for lado, val in (("top", arriba), ("bottom", abajo),
                      ("left", izq), ("right", der)):
        el = OxmlElement(f"w:{lado}")
        el.set(qn("w:w"), str(val))
        el.set(qn("w:type"), "dxa")
        mar.append(el)
    tbl_pr.append(mar)


def _anchos(tabla, columnas, ancho_total) -> None:
    """Reparte el ancho de la página segun el `ancho` que declara cada columna.

    ⚠️ Antes la tabla iba con `autofit`, y Word repartía por el CONTENIDO: una
    columna con «TOTAL RENT AND MANAGEMENT FEES» se comía el espacio de los
    números. Los `ancho` ya viajaban en el payload y no los miraba nadie.
    """
    pesos = [max(1, int(c.get("ancho") or 12)) for c in columnas]
    total = sum(pesos)
    tabla.autofit = False
    for fila in tabla.rows:
        for i, celda in enumerate(fila.cells):
            celda.width = Emu(int(ancho_total * pesos[i] / total))


def _borde_caja(tabla) -> None:
    """Marco fino alrededor de una tabla de una celda (el de comentarios)."""
    tbl_pr = tabla._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for lado in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{lado}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "6")
        el.set(qn("w:color"), "C7CCD1")
        borders.append(el)
    tbl_pr.append(borders)


def _campo(parrafo, instruccion: str) -> None:
    """Un campo de Word (`PAGE`, `NUMPAGES`). Sin esto el pie diría «1» en todas."""
    r = parrafo.add_run()
    ini = OxmlElement("w:fldChar"); ini.set(qn("w:fldCharType"), "begin")
    txt = OxmlElement("w:instrText"); txt.set(qn("xml:space"), "preserve")
    txt.text = f" {instruccion} "
    fin = OxmlElement("w:fldChar"); fin.set(qn("w:fldCharType"), "end")
    r._r.append(ini); r._r.append(txt); r._r.append(fin)


def _apaisar(seccion) -> None:
    """⚠️ No alcanza con cambiar la orientación: Word no gira la hoja solo. Hay
    que intercambiar ancho y alto, o queda una carta vertical rotulada como
    apaisada y el cuadro sigue sin entrar."""
    ancho, alto = seccion.page_width, seccion.page_height
    seccion.orientation = WD_ORIENT.LANDSCAPE
    seccion.page_width, seccion.page_height = max(ancho, alto), min(ancho, alto)


def _enderezar(seccion) -> None:
    ancho, alto = seccion.page_width, seccion.page_height
    seccion.orientation = WD_ORIENT.PORTRAIT
    seccion.page_width, seccion.page_height = min(ancho, alto), max(ancho, alto)


def _margenes(seccion) -> None:
    for lado in ("top", "bottom"):
        setattr(seccion, f"{lado}_margin", Cm(1.8))
    for lado in ("left", "right"):
        setattr(seccion, f"{lado}_margin", Cm(1.6))


def _pie(seccion, propiedad: str, periodo: str) -> None:
    p = seccion.footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"{propiedad} · {periodo} · página ")
    r.font.size = Pt(8); r.font.color.rgb = GRIS
    _campo(p, "PAGE")
    r2 = p.add_run(" de ")
    r2.font.size = Pt(8); r2.font.color.rgb = GRIS
    _campo(p, "NUMPAGES")
    for run in p.runs:
        run.font.size = Pt(8)
        run.font.color.rgb = GRIS


# ─── El documento ─────────────────────────────────────────────────────────────
def _portada(doc, propiedad: str, titulo: str, periodo: str,
             versiones: str, cuadros: list[dict]) -> None:
    for _ in range(4):
        doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(propiedad.upper())
    r.font.size = Pt(13); r.font.bold = True; r.font.color.rgb = GRIS

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(titulo)
    r.font.size = Pt(26); r.font.bold = True; r.font.color.rgb = MARCA

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(periodo)
    r.font.size = Pt(15); r.font.color.rgb = NEGRO

    if versiones:
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(versiones)
        r.font.size = Pt(11); r.font.color.rgb = GRIS

    for _ in range(2):
        doc.add_paragraph()

    # El índice: qué cuadros trae el documento. Sale de los cuadros que llegaron,
    # así que refleja exactamente lo que está activo en la vista.
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("CONTENIDO")
    r.font.size = Pt(9); r.font.bold = True; r.font.color.rgb = GRIS
    for i, c in enumerate(cuadros, 1):
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(f"{i}.  {(c.get('titulo') or 'Cuadro').strip()}")
        r.font.size = Pt(10.5); r.font.color.rgb = NEGRO

    for _ in range(3):
        doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"Generado el {date.today():%d/%m/%Y}")
    r.font.size = Pt(9); r.font.color.rgb = GRIS


def _tabla(doc, cuadro: dict) -> None:
    columnas = cuadro.get("columnas") or []
    filas = cuadro.get("filas") or []
    if not columnas:
        return

    tabla = doc.add_table(rows=1, cols=len(columnas))
    tabla.alignment = WD_TABLE_ALIGNMENT.CENTER
    _sin_rejilla(tabla)
    _apretar(tabla)

    # ── Encabezado ────────────────────────────────────────────────────────────
    encabezado = tabla.rows[0]
    encabezado.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
    for i, col in enumerate(columnas):
        celda = encabezado.cells[i]
        p = _vaciar(celda)
        p.alignment = (WD_ALIGN_PARAGRAPH.LEFT if i == 0
                       else WD_ALIGN_PARAGRAPH.RIGHT)
        r = p.add_run(str(col.get("label") or ""))
        r.font.size = Pt(7.5)
        r.font.bold = True
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        _sombra(celda, "1B3A5C")
        # La banda del encabezado ya separa; la regla la cierra por abajo.
        _regla(celda, "bottom", 8, "1B3A5C")
    # Que el encabezado se repita si la tabla parte de página.
    trPr = encabezado._tr.get_or_add_trPr()
    th = OxmlElement("w:tblHeader"); th.set(qn("w:val"), "true")
    trPr.append(th)

    # ── Cuerpo ────────────────────────────────────────────────────────────────
    for n, fila in enumerate(filas):
        celdas = tabla.add_row().cells
        es_total = bool(fila.get("es_total"))
        formato_fila = fila.get("formato")
        nivel = int(fila.get("nivel") or 0)

        c0 = celdas[0]
        p0 = _vaciar(c0)
        r0 = p0.add_run("    " * nivel + str(fila.get("label") or ""))
        r0.font.size = Pt(7.5)
        r0.font.bold = es_total
        r0.font.color.rgb = NEGRO

        valores = fila.get("valores") or []
        for i in range(1, len(columnas)):
            v = valores[i - 1] if i - 1 < len(valores) else None
            formato = formato_fila or (columnas[i].get("formato") or "usd")
            celda = celdas[i]
            p = _vaciar(celda)
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            r = p.add_run(_texto(v, formato))
            r.font.size = Pt(7.5)
            r.font.bold = es_total
            r.font.color.rgb = ROJO if _negativo(v) else NEGRO

        # ── Cómo se separa una fila de la siguiente ───────────────────────
        #
        # Un TOTAL lleva regla arriba —la convención del estado impreso: la
        # línea dice «acá se cierra algo»— y fondo suave. Las demás no llevan
        # nada: la cebra sola alcanza, y una regla bajo cada fila devuelve la
        # rejilla que se acaba de sacar.
        if es_total:
            for celda in celdas:
                _sombra(celda, "E8EDF2")
                _regla(celda, "top", 8, "1B3A5C")
        elif n % 2:
            for celda in celdas:
                _sombra(celda, "F7F8FA")

    # El ancho, al final: hay que repartirlo sobre las filas ya creadas.
    seccion = doc.sections[-1]
    _anchos(tabla, columnas,
            seccion.page_width - seccion.left_margin - seccion.right_margin)


def _comentarios(doc, notas: list[str] | None = None) -> None:
    """El recuadro para comentar, con lo ya escrito adentro.

    Owner, 2026-09-03: *«una vez que se impriman en Word, estas notas aparezcan
    en el box editable»*.

    ⚠️ Las notas van DENTRO del recuadro y no en un párrafo aparte: el recuadro
    es el lugar donde se comenta, y separar lo escrito de donde se escribe haría
    que en la reunión se comente dos veces lo mismo — una en el papel y otra en
    la app.

    Los renglones en blanco se agregan DESPUÉS de las notas, para que siempre
    quede espacio para escribir a mano encima de lo que ya está.
    """
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run("COMENTARIOS")
    r.font.size = Pt(8.5); r.font.bold = True; r.font.color.rgb = GRIS

    caja = doc.add_table(rows=1, cols=1)
    _borde_caja(caja)
    celda = caja.rows[0].cells[0]
    celda.text = ""
    primero = True

    for nota in (notas or []):
        texto = str(nota or "").strip()
        if not texto:
            continue
        p = celda.paragraphs[0] if primero else celda.add_paragraph()
        primero = False
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(texto)
        r.font.size = Pt(9)
        r.font.color.rgb = NEGRO

    for _ in range(RENGLONES_COMENTARIO):
        p = celda.paragraphs[0] if primero else celda.add_paragraph()
        primero = False
        p.paragraph_format.space_after = Pt(9)
        p.add_run("")


def build_cierre_docx(cuadros: list[dict], propiedad: str, titulo: str,
                      periodo: str, versiones: str = "") -> bytes:
    """El documento completo. Devuelve los bytes del `.docx`."""
    doc = Document()

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10)

    portada = doc.sections[0]
    _margenes(portada)
    _enderezar(portada)
    _pie(portada, propiedad, periodo)
    _portada(doc, propiedad, titulo, periodo, versiones, cuadros)

    for cuadro in cuadros:
        ancho = len(cuadro.get("columnas") or []) >= COLUMNAS_APAISADO
        seccion = doc.add_section(WD_SECTION.NEW_PAGE)
        _margenes(seccion)
        _apaisar(seccion) if ancho else _enderezar(seccion)
        # ⚠️ Cada sección hereda el pie de la anterior sólo si se lo pide; con
        # `is_linked_to_previous` en False habría que rearmarlo en cada una.
        seccion.footer.is_linked_to_previous = True

        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run((cuadro.get("titulo") or "Cuadro").strip())
        r.font.size = Pt(15); r.font.bold = True; r.font.color.rgb = MARCA

        sub = (cuadro.get("subtitulo") or "").strip()
        if sub:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(8)
            r = p.add_run(sub)
            r.font.size = Pt(9.5); r.font.color.rgb = GRIS

        _tabla(doc, cuadro)
        _comentarios(doc, cuadro.get("comentarios"))

    salida = io.BytesIO()
    doc.save(salida)
    return salida.getvalue()
